"""
Extracts plain text from the NZ Building Code reference .docx
and writes it to nzbc_knowledge.txt for the bot to use.

Run once: python extract_knowledge.py
"""

import sys
import subprocess
import importlib.util

def ensure(pkg, import_as=None):
    name = import_as or pkg
    if importlib.util.find_spec(name) is None:
        subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

ensure("python-docx", "docx")

from docx import Document  # noqa: E402
from pathlib import Path   # noqa: E402

DOCX_PATH = Path(__file__).parent / "NZ-Building-Code-Reference.docx"
OUT_PATH = Path(__file__).parent / "nzbc_knowledge.txt"


def extract(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    # Also extract table cell text
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_parts.append(cell_text)
            if row_parts:
                lines.append(" | ".join(row_parts))
    return "\n".join(lines)


if __name__ == "__main__":
    if not DOCX_PATH.exists():
        print(f"ERROR: {DOCX_PATH} not found.")
        print("Copy NZ-Building-Code-Reference.docx into this directory first.")
        sys.exit(1)

    print(f"Extracting from {DOCX_PATH}...")
    text = extract(DOCX_PATH)
    OUT_PATH.write_text(text, encoding="utf-8")
    print(f"Written {len(text):,} characters to {OUT_PATH}")
